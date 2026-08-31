from io import BytesIO

import pymupdf
from docx import Document as WordDocument

from app.core.database import SessionLocal
from app.models.document import Document, DocumentCategory, DocumentStatus
from app.services.text_chunker import ExtractedBlock, ExtractedPage, TextChunker


def make_pdf_bytes(*pages: str) -> bytes:
    pdf = pymupdf.open()
    for text in pages:
        page = pdf.new_page()
        page.insert_text((72, 72), text)
    result = pdf.tobytes()
    pdf.close()
    return result


def make_docx_bytes() -> bytes:
    document = WordDocument()
    document.add_heading("Customer Concentration", level=1)
    document.add_paragraph("Customer A represents 31% of revenue.")
    document.add_paragraph("Customer B represents 18%.")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_health(client):
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_list_deals_includes_seed_deal(client):
    response = client.get("/api/deals")
    assert response.status_code == 200
    assert response.json()[0]["name"] == "ABC Hydraulic Systems"


def test_get_seed_deal(client):
    response = client.get("/api/deals/ABC-HYD-001")
    assert response.status_code == 200
    assert response.json()["revenue"] == "120000000.00"
    assert response.json()["deal_stage"] == "Due Diligence"


def test_seed_deal_has_no_documents(client):
    response = client.get("/api/deals/ABC-HYD-001/documents")
    assert response.status_code == 200
    assert response.json() == []


def test_document_is_scoped_to_its_deal(client):
    with SessionLocal() as db:
        db.add(
            Document(
                document_id="DOC-001",
                deal_id="ABC-HYD-001",
                file_name="financials.pdf",
                category=DocumentCategory.FINANCIAL,
                file_path="data/documents/financials.pdf",
                status=DocumentStatus.PENDING,
            )
        )
        db.commit()

    response = client.get("/api/deals/ABC-HYD-001/documents")
    assert response.status_code == 200
    assert [document["document_id"] for document in response.json()] == ["DOC-001"]


def test_pdf_upload_extracts_pages_chunks_and_downloads(client):
    pdf = make_pdf_bytes(
        "Customer Concentration\nCustomer A represents 31% of revenue.",
        "Customer B represents 18% of revenue.",
    )
    response = client.post(
        "/api/deals/ABC-HYD-001/documents",
        data={"category": "Commercial"},
        files={"file": ("Customer Analysis.pdf", pdf, "application/pdf")},
    )
    assert response.status_code == 201
    document = response.json()
    assert document["category"] == "Commercial"
    assert document["status"] == "Processed"
    assert document["page_count"] == 2
    assert document["chunk_count"] == 2

    listed = client.get("/api/deals/ABC-HYD-001/documents")
    assert [item["document_id"] for item in listed.json()] == [document["document_id"]]

    chunks = client.get(f"/api/documents/{document['document_id']}/chunks")
    assert [chunk["page_number"] for chunk in chunks.json()] == [1, 2]
    assert "Customer A represents 31% of revenue." in chunks.json()[0]["chunk_text"]

    download = client.get(f"/api/documents/{document['document_id']}/download")
    assert download.status_code == 200
    assert download.content == pdf
    assert download.headers["content-type"] == "application/pdf"


def test_docx_extraction_preserves_heading_and_text(client):
    response = client.post(
        "/api/deals/ABC-HYD-001/documents",
        data={"category": "Commercial"},
        files={"file": ("Customer Analysis.docx", make_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 201
    document = response.json()
    assert document["status"] == "Processed"
    assert document["page_count"] == 1
    chunk = client.get(f"/api/documents/{document['document_id']}/chunks").json()[0]
    assert chunk["page_number"] is None
    assert chunk["section"] == "Customer Concentration"
    assert "Customer A represents 31% of revenue." in chunk["chunk_text"]


def test_chunker_is_deterministic_and_overlaps():
    words = " ".join(f"word{index}" for index in range(220))
    chunks = TextChunker(100, 20).chunk([ExtractedPage(1, [ExtractedBlock(words)])])
    assert len(chunks) == 3
    assert all(chunk.token_count <= 100 for chunk in chunks)
    assert chunks[0].chunk_text.split()[-20:] == chunks[1].chunk_text.split()[:20]


def test_rejects_unsupported_document_type(client):
    response = client.post(
        "/api/deals/ABC-HYD-001/documents",
        data={"category": "Legal"},
        files={"file": ("script.exe", b"not allowed", "application/octet-stream")},
    )
    assert response.status_code == 400
    assert "Unsupported file type" in response.json()["detail"]


def test_xlsx_processing_failure_is_recorded(client):
    response = client.post(
        "/api/deals/ABC-HYD-001/documents",
        data={"category": "Financial"},
        files={"file": ("financials.xlsx", b"not a real workbook", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert response.status_code == 201
    document = response.json()
    assert document["status"] == "Failed"
    assert "not implemented" in document["processing_error"]


def test_reprocessing_replaces_chunks(client):
    response = client.post(
        "/api/deals/ABC-HYD-001/documents",
        data={"category": "Commercial"},
        files={"file": ("Customers.pdf", make_pdf_bytes("Customer A represents 31% of revenue."), "application/pdf")},
    )
    document_id = response.json()["document_id"]
    first_chunks = client.get(f"/api/documents/{document_id}/chunks").json()
    reprocessed = client.post(f"/api/documents/{document_id}/process")
    second_chunks = client.get(f"/api/documents/{document_id}/chunks").json()
    assert reprocessed.json()["status"] == "Processed"
    assert len(second_chunks) == len(first_chunks)
    assert [chunk["chunk_index"] for chunk in second_chunks] == list(range(len(second_chunks)))


def test_documents_are_isolated_between_deals(client):
    client.post(
        "/api/deals",
        json={
            "deal_id": "SECOND-001", "name": "Second Deal", "company_name": "Second Co",
            "industry": "Software", "geography": "Europe", "revenue": 1, "ebitda": 1,
            "deal_stage": "Screening",
        },
    )
    upload = client.post(
        "/api/deals/SECOND-001/documents",
        data={"category": "Financial"},
        files={"file": ("second.pdf", make_pdf_bytes("Second deal only"), "application/pdf")},
    )
    assert upload.status_code == 201
    assert client.get("/api/deals/ABC-HYD-001/documents").json() == []
    chunks = client.get(f"/api/documents/{upload.json()['document_id']}/chunks").json()
    assert {chunk["deal_id"] for chunk in chunks} == {"SECOND-001"}
