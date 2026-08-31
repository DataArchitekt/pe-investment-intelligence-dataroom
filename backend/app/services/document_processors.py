from pathlib import Path

import pymupdf
from docx import Document as WordDocument

from app.integrations.contracts import DocumentProcessor
from app.services.text_chunker import ExtractedBlock, ExtractedPage, clean_text


class ProcessingError(ValueError):
    pass


class PDFDocumentProcessor(DocumentProcessor):
    def extract(self, source: Path) -> list[ExtractedPage]:
        try:
            pdf = pymupdf.open(source)
        except Exception as error:
            raise ProcessingError("The PDF could not be opened.") from error
        try:
            if pdf.page_count == 0:
                raise ProcessingError("The PDF has no pages.")
            pages = []
            for index, page in enumerate(pdf):
                text = clean_text(page.get_text("text"))
                if text:
                    pages.append(ExtractedPage(index + 1, [ExtractedBlock(text)]))
            if not pages:
                raise ProcessingError("No extractable text was found in the PDF.")
            return pages
        finally:
            pdf.close()


class DOCXDocumentProcessor(DocumentProcessor):
    def extract(self, source: Path) -> list[ExtractedPage]:
        try:
            document = WordDocument(source)
        except Exception as error:
            raise ProcessingError("The DOCX file could not be opened.") from error

        blocks: list[ExtractedBlock] = []
        current_section: str | None = None
        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            if not text:
                continue
            if paragraph.style and paragraph.style.name.startswith("Heading"):
                current_section = text
                continue
            blocks.append(ExtractedBlock(text, current_section))
        for table in document.tables:
            for row in table.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    blocks.append(ExtractedBlock(row_text, current_section))
        if not blocks:
            raise ProcessingError("No extractable text was found in the DOCX file.")
        return [ExtractedPage(None, blocks)]


class TXTDocumentProcessor(DocumentProcessor):
    def extract(self, source: Path) -> list[ExtractedPage]:
        try:
            text = clean_text(source.read_text(encoding="utf-8", errors="replace"))
        except OSError as error:
            raise ProcessingError("The text file could not be read.") from error
        if not text:
            raise ProcessingError("The text file is empty.")
        return [ExtractedPage(None, [ExtractedBlock(text)])]


def processor_for(source: Path) -> DocumentProcessor:
    processors: dict[str, DocumentProcessor] = {
        ".pdf": PDFDocumentProcessor(),
        ".docx": DOCXDocumentProcessor(),
        ".txt": TXTDocumentProcessor(),
    }
    try:
        return processors[source.suffix.lower()]
    except KeyError as error:
        raise ProcessingError(f"Processing is not implemented for {source.suffix.upper() or 'this file type'}.") from error
